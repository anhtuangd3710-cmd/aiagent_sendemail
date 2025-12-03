"""
Email AI Agent - Flask Web Application
Modern UI for email automation with Azure OpenAI or Google Gemini
With Realtime WebSocket support, Authentication, and Auto-start Monitor
"""
from flask import Flask, render_template, request, jsonify, Response, make_response, render_template_string, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import json
import threading
import queue
import atexit
import os
import io
from datetime import datetime

from services.email_service import EmailService
from services.email_monitor import EmailMonitor, ManualResponseProcessor
from services.auth_service import AuthService, login_required, admin_required
from config.settings import SENDER_EMAIL, AI_PROVIDER, AUTO_START_MONITOR

# Import Database Service based on DATABASE_URL
DATABASE_URL = os.getenv('DATABASE_URL')
if DATABASE_URL:
    from services.database_postgres import DatabaseServicePostgres as DatabaseService
    print("🗄️ Using PostgreSQL (Neon Database)")
else:
    from services.database import DatabaseService
    print("🗄️ Using SQLite")

# Import AI services based on provider
if AI_PROVIDER == "gemini":
    from services.ai_agent_gemini import AIAgentGemini as AIAgent
    from services.cv_evaluator_gemini import CVEvaluatorGemini as CVEvaluator
    print("🤖 Using Google Gemini AI")
else:
    from services.ai_agent import AIAgent
    from services.cv_evaluator import CVEvaluator
    print("🤖 Using Azure OpenAI")


def get_ai_agent_for_user(user_settings: dict):
    """
    Get AI agent with user's API key if configured, otherwise use default.
    If user has their own Gemini API key, use it for unlimited usage.
    """
    import google.generativeai as genai
    from config.settings import GEMINI_MODEL
    
    user_gemini_key = user_settings.get('gemini_api_key') if user_settings else None
    
    if user_gemini_key:
        # Create a custom AI agent with user's API key
        class UserAIAgent:
            def __init__(self, api_key):
                genai.configure(api_key=api_key)
                self.model = genai.GenerativeModel(
                    model_name=GEMINI_MODEL,
                    generation_config={
                        "temperature": 0.7,
                        "top_p": 0.95,
                        "top_k": 40,
                        "max_output_tokens": 8192,
                    }
                )
                # Copy methods from the default ai_agent
                self._extract_json = ai_agent._extract_json
                self.generate_email = self._generate_email_method
                self.analyze_response = ai_agent.analyze_response
                self.generate_notification_email = ai_agent.generate_notification_email
            
            def _generate_email_method(self, *args, **kwargs):
                """Generate email using user's model"""
                # Re-use the base agent's logic but with user's model
                return ai_agent.generate_email(*args, **kwargs)
        
        return UserAIAgent(user_gemini_key)
    else:
        # Use default AI agent (with system API key - limited free usage)
        return ai_agent


def get_cv_evaluator_for_user(user_settings: dict):
    """
    Get CV evaluator with user's API key if configured, otherwise use default.
    """
    import google.generativeai as genai
    from config.settings import GEMINI_MODEL
    
    user_gemini_key = user_settings.get('gemini_api_key') if user_settings else None
    
    if user_gemini_key:
        # Configure with user's API key and return default evaluator
        # Note: This reconfigures genai globally for this request
        genai.configure(api_key=user_gemini_key)
    
    return cv_evaluator


app = Flask(__name__)
app.secret_key = 'email-agent-secret-key-change-in-production'
CORS(app, supports_credentials=True)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading')

# ==================== Performance Optimization ====================

# Server-side cache for reducing database calls
from functools import lru_cache
import time

class ServerCache:
    """Simple TTL cache for server-side caching"""
    def __init__(self):
        self._cache = {}
        self._timestamps = {}
        self._ttl = {
            'emails': 10,       # 10 seconds for emails
            'cv_list': 30,      # 30 seconds for CV list
            'stats': 30,        # 30 seconds for stats
            'settings': 60,     # 1 minute for user settings
        }
    
    def get(self, key):
        if key in self._cache:
            ttl = self._ttl.get(key.split(':')[0], 30)
            if time.time() - self._timestamps.get(key, 0) < ttl:
                return self._cache[key]
            else:
                del self._cache[key]
                del self._timestamps[key]
        return None
    
    def set(self, key, value):
        self._cache[key] = value
        self._timestamps[key] = time.time()
    
    def invalidate(self, pattern=None):
        if pattern:
            keys_to_delete = [k for k in self._cache if pattern in k]
            for k in keys_to_delete:
                del self._cache[k]
                del self._timestamps[k]
        else:
            self._cache.clear()
            self._timestamps.clear()

server_cache = ServerCache()

# Response compression middleware
@app.after_request
def add_cache_headers(response):
    """Add cache headers and compression hints"""
    # For API responses, add cache control
    if request.path.startswith('/api/'):
        # Short cache for dynamic data
        if 'emails' in request.path or 'cv' in request.path:
            response.headers['Cache-Control'] = 'private, max-age=5'
        elif 'stats' in request.path:
            response.headers['Cache-Control'] = 'private, max-age=30'
        else:
            response.headers['Cache-Control'] = 'no-cache'
    
    # Add compression hint
    response.headers['Vary'] = 'Accept-Encoding'
    
    return response

# Initialize services
email_service = EmailService()
ai_agent = AIAgent()
database = DatabaseService()
cv_evaluator = CVEvaluator()
auth_service = AuthService(database)

# Store auth_service in app config for decorators
app.config['auth_service'] = auth_service

# Event queue for real-time updates (kept for backward compatibility)
event_queue = queue.Queue()

# Monitor instance - will be auto-started
monitor = None
monitor_auto_started = False

# Connected clients count
connected_clients = 0


def notification_callback(email_record, response, analysis):
    """Callback when a response is received - emit via WebSocket"""
    # Invalidate email cache when new response received
    server_cache.invalidate('emails')
    
    event_data = {
        "type": "response_received",
        "email_id": email_record['id'],
        "recipient_name": email_record['recipient_name'],
        "recipient_email": email_record['recipient_email'],
        "analysis": analysis,
        "timestamp": datetime.now().isoformat()
    }
    
    # Emit to all connected WebSocket clients
    socketio.emit('response_received', event_data)
    
    # Also put in queue for SSE fallback
    event_queue.put(event_data)


def cv_notification_callback(cv_evaluation, email_sent=False):
    """Callback for CV evaluation events"""
    event_data = {
        "type": "cv_evaluated",
        "cv_id": cv_evaluation.get('id'),
        "candidate_name": cv_evaluation.get('candidate_name'),
        "overall_score": cv_evaluation.get('overall_score'),
        "is_qualified": cv_evaluation.get('is_qualified'),
        "email_sent": email_sent,
        "timestamp": datetime.now().isoformat()
    }
    socketio.emit('cv_evaluated', event_data)


def auto_start_monitor():
    """Auto-start the email monitor on application startup"""
    global monitor, monitor_auto_started
    
    if monitor_auto_started:
        return
    
    try:
        monitor = EmailMonitor(
            email_service,
            ai_agent,
            database,
            notification_callback=notification_callback
        )
        monitor.start()
        monitor_auto_started = True
        print("✅ Email monitor auto-started successfully")
    except Exception as e:
        print(f"⚠️ Failed to auto-start monitor: {e}")


def stop_monitor_on_exit():
    """Stop monitor when application exits"""
    global monitor
    if monitor:
        try:
            monitor.stop()
            print("Monitor stopped on exit")
        except:
            pass


# Register cleanup
atexit.register(stop_monitor_on_exit)


# WebSocket Events
@socketio.on('connect')
def handle_connect():
    """Handle client connection"""
    global connected_clients
    connected_clients += 1
    
    # Auto-start monitor on first client connection
    if not monitor_auto_started:
        auto_start_monitor()
    
    emit('connected', {
        'status': 'connected', 
        'clients': connected_clients,
        'monitor_running': monitor is not None and monitor._running if monitor else False
    })
    print(f"Client connected. Total clients: {connected_clients}")


@socketio.on('disconnect')
def handle_disconnect():
    """Handle client disconnection"""
    global connected_clients
    connected_clients -= 1
    print(f"Client disconnected. Total clients: {connected_clients}")


@socketio.on('check_responses_now')
def handle_check_responses():
    """Handle immediate response check request from client"""
    try:
        temp_monitor = EmailMonitor(
            email_service,
            ai_agent,
            database,
            notification_callback=notification_callback
        )
        temp_monitor.check_responses()
        emit('check_complete', {'success': True, 'message': 'Response check completed'})
    except Exception as e:
        emit('check_complete', {'success': False, 'error': str(e)})


# ==================== Authentication Routes ====================

@app.route('/api/auth/register', methods=['POST'])
def register():
    """Register a new user"""
    try:
        data = request.json
        result = auth_service.register(
            username=data.get('username'),
            email=data.get('email'),
            password=data.get('password'),
            full_name=data.get('full_name')
        )
        
        if result['success']:
            return jsonify(result), 201
        return jsonify(result), 400
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auth/login', methods=['POST'])
def login():
    """Login user"""
    try:
        data = request.json
        ip_address = request.remote_addr
        user_agent = request.headers.get('User-Agent')
        
        result = auth_service.login(
            username_or_email=data.get('username') or data.get('email'),
            password=data.get('password'),
            ip_address=ip_address,
            user_agent=user_agent
        )
        
        if result['success']:
            response = make_response(jsonify(result))
            # Set cookie for browser
            response.set_cookie(
                'auth_token', 
                result['token'],
                httponly=True,
                secure=False,  # Set True in production with HTTPS
                samesite='Lax',
                max_age=7*24*60*60  # 7 days
            )
            return response
        return jsonify(result), 401
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auth/logout', methods=['POST'])
def logout():
    """Logout user"""
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            token = request.cookies.get('auth_token')
        
        if token:
            auth_service.logout(token)
        
        response = make_response(jsonify({"success": True, "message": "Đăng xuất thành công"}))
        response.delete_cookie('auth_token')
        return response
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auth/me', methods=['GET'])
def get_current_user():
    """Get current user info"""
    try:
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            token = request.cookies.get('auth_token')
        
        if not token:
            return jsonify({"success": False, "error": "Chưa đăng nhập"}), 401
        
        user = auth_service.validate_token(token)
        if not user:
            return jsonify({"success": False, "error": "Phiên đăng nhập hết hạn"}), 401
        
        return jsonify({"success": True, "user": user})
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auth/change-password', methods=['POST'])
@login_required
def change_password():
    """Change user password"""
    try:
        data = request.json
        result = auth_service.change_password(
            user_id=request.current_user['id'],
            old_password=data.get('old_password'),
            new_password=data.get('new_password')
        )
        
        if result['success']:
            return jsonify(result)
        return jsonify(result), 400
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auth/profile', methods=['PUT'])
@login_required
def update_profile():
    """Update user profile"""
    try:
        data = request.json
        result = auth_service.update_profile(
            user_id=request.current_user['id'],
            full_name=data.get('full_name'),
            email=data.get('email')
        )
        
        if result['success']:
            return jsonify(result)
        return jsonify(result), 400
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/admin/users', methods=['GET'])
@admin_required
def get_all_users():
    """Get all users (admin only)"""
    try:
        users = auth_service.get_all_users()
        return jsonify({"success": True, "users": users})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== User Settings API ====================

@app.route('/api/user/settings', methods=['GET'])
@login_required
def get_user_settings():
    """Get current user's settings"""
    try:
        settings = database.get_user_settings(request.current_user['id'])
        if settings:
            # Mask sensitive data
            masked_settings = dict(settings)
            if masked_settings.get('azure_openai_api_key'):
                masked_settings['azure_openai_api_key'] = '***' + masked_settings['azure_openai_api_key'][-4:]
            if masked_settings.get('gemini_api_key'):
                masked_settings['gemini_api_key'] = '***' + masked_settings['gemini_api_key'][-4:]
            if masked_settings.get('sender_password'):
                masked_settings['sender_password'] = '********'
            return jsonify({"success": True, "settings": masked_settings})
        return jsonify({"success": True, "settings": None})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/user/settings', methods=['POST', 'PUT'])
@login_required
def save_user_settings():
    """Save user settings"""
    try:
        data = request.json
        settings_id = database.save_user_settings(request.current_user['id'], data)
        return jsonify({
            "success": True,
            "message": "Cài đặt đã được lưu",
            "settings_id": settings_id
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/user/settings/full', methods=['GET'])
@login_required
def get_user_settings_full():
    """Get full settings (unmasked) for internal use"""
    try:
        settings = database.get_user_settings(request.current_user['id'])
        return jsonify({"success": True, "settings": settings})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== Main Routes ====================

@app.route('/')
def index():
    """Main page - requires login"""
    return render_template('index.html')


@app.route('/login')
def login_page():
    """Login page"""
    return render_template('login.html')


@app.route('/api/send-email', methods=['POST'])
@login_required
def send_email():
    """Send an AI-generated email with optional attachments"""
    try:
        # Check if this is a multipart form (with attachments) or JSON
        if request.content_type and 'multipart/form-data' in request.content_type:
            # Handle form data with attachments
            data = request.form.to_dict()
            files = request.files.getlist('attachments')
        else:
            data = request.json
            files = []
        
        # Get user settings for email
        user_settings = database.get_user_settings(request.current_user['id'])
        
        # Check if email is configured - REQUIRED for sending
        if not user_settings or not user_settings.get('sender_email') or not user_settings.get('sender_password'):
            return jsonify({
                "success": False,
                "error": "Bạn chưa cấu hình Email gửi. Vui lòng vào mục 'Hồ sơ & API Keys' để cấu hình Email và App Password trước khi gửi email.",
                "error_code": "EMAIL_NOT_CONFIGURED"
            }), 400
        
        # Create email service with user's settings
        user_email_service = EmailService()
        user_email_service.sender_email = user_settings.get('sender_email')
        user_email_service.sender_password = user_settings.get('sender_password')
        user_email_service.smtp_host = user_settings.get('email_host', 'smtp.gmail.com')
        user_email_service.smtp_port = int(user_settings.get('email_port', 587))
        sender_email_config = user_settings.get('sender_email')
        
        sender_name = data.get('sender_name')
        recipient_name = data.get('recipient_name')
        recipient_email = data.get('recipient_email')
        purpose = data.get('purpose')
        tone = data.get('tone', 'professional')
        language = data.get('language', 'vi')
        additional_context = data.get('additional_context')
        notification_email = data.get('notification_email') or sender_email_config
        
        # Check if custom subject/body provided (from preview)
        custom_subject = data.get('custom_subject')
        custom_body = data.get('custom_body')
        
        # Validate required fields
        if not all([sender_name, recipient_name, recipient_email, purpose]):
            return jsonify({
                "success": False,
                "error": "Missing required fields"
            }), 400
        
        # Generate email using AI or use custom content
        if custom_subject and custom_body:
            generated_email = {
                'subject': custom_subject,
                'body': custom_body
            }
        else:
            # Use user's Gemini API key if configured
            user_ai_agent = get_ai_agent_for_user(user_settings)
            generated_email = user_ai_agent.generate_email(
                sender_name=sender_name,
                recipient_name=recipient_name,
                recipient_email=recipient_email,
                purpose=purpose,
                tone=tone,
                language=language,
                additional_context=additional_context
            )
        
        # Process attachments
        attachments = []
        for file in files:
            if file and file.filename:
                # Read file content
                content = file.read()
                # Get content type
                content_type = file.content_type or 'application/octet-stream'
                
                attachments.append({
                    'filename': file.filename,
                    'content': content,
                    'content_type': content_type
                })
        
        # Send the email with attachments using user's email service
        success = user_email_service.send_email(
            recipient_email=recipient_email,
            subject=generated_email['subject'],
            body=generated_email['body'],
            attachments=attachments if attachments else None
        )
        
        if success:
            # Invalidate email cache for this user
            user_id = request.current_user['id']
            server_cache.invalidate(f'emails:{user_id}')
            server_cache.invalidate(f'stats:{user_id}')
            
            # Save to database
            email_id = database.save_sent_email(
                sender_name=sender_name,
                sender_email=notification_email,
                recipient_name=recipient_name,
                recipient_email=recipient_email,
                subject=generated_email['subject'],
                body=generated_email['body'],
                purpose=purpose,
                user_id=user_id
            )
            
            # Auto-start monitor after sending email
            global monitor, monitor_auto_started
            if not monitor or not monitor._running:
                try:
                    # Create monitor with user's email settings if available
                    if user_settings and user_settings.get('sender_email') and user_settings.get('sender_password'):
                        user_monitor_email_service = EmailService()
                        user_monitor_email_service.sender_email = user_settings.get('sender_email')
                        user_monitor_email_service.sender_password = user_settings.get('sender_password')
                        user_monitor_email_service.imap_host = user_settings.get('imap_host', 'imap.gmail.com')
                        user_monitor_email_service.imap_port = int(user_settings.get('imap_port', 993))
                        
                        monitor = EmailMonitor(
                            user_monitor_email_service,
                            ai_agent,
                            database,
                            notification_callback=notification_callback
                        )
                    else:
                        monitor = EmailMonitor(
                            email_service,
                            ai_agent,
                            database,
                            notification_callback=notification_callback
                        )
                    
                    monitor.start()
                    monitor_auto_started = True
                    print("✅ Email monitor auto-started after sending email")
                except Exception as e:
                    print(f"⚠️ Failed to auto-start monitor: {e}")
            
            return jsonify({
                "success": True,
                "email_id": email_id,
                "subject": generated_email['subject'],
                "body": generated_email['body'],
                "message": "Email sent successfully!",
                "monitor_started": monitor is not None and monitor._running if monitor else False
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to send email. Please check your email configuration."
            }), 500
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/preview-email', methods=['POST'])
@login_required
def preview_email():
    """Preview AI-generated email without sending"""
    try:
        data = request.json
        
        sender_name = data.get('sender_name')
        recipient_name = data.get('recipient_name')
        recipient_email = data.get('recipient_email')
        purpose = data.get('purpose')
        tone = data.get('tone', 'professional')
        language = data.get('language', 'vi')
        additional_context = data.get('additional_context')
        
        if not all([sender_name, recipient_name, recipient_email, purpose]):
            return jsonify({
                "success": False,
                "error": "Missing required fields"
            }), 400
        
        # Get user settings to check for custom API key
        user_id = request.current_user.get('id')
        user_settings = database.get_user_settings(user_id) if user_id else None
        
        # Use user's AI agent if they have custom API key, otherwise use default
        user_ai_agent = get_ai_agent_for_user(user_settings)
        
        # Generate email using AI
        generated_email = user_ai_agent.generate_email(
            sender_name=sender_name,
            recipient_name=recipient_name,
            recipient_email=recipient_email,
            purpose=purpose,
            tone=tone,
            language=language,
            additional_context=additional_context
        )
        
        return jsonify({
            "success": True,
            "subject": generated_email['subject'],
            "body": generated_email['body']
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails', methods=['GET'])
@login_required
def get_emails():
    """Get all tracked emails with server-side caching"""
    try:
        user_id = request.current_user['id']
        cache_key = f'emails:{user_id}'
        
        # Check server cache first
        cached = server_cache.get(cache_key)
        if cached:
            return jsonify({
                "success": True,
                "emails": cached,
                "cached": True
            })
        
        emails = database.get_all_emails(user_id)
        
        # Parse analysis JSON for each email
        for email in emails:
            if email.get('analysis'):
                try:
                    email['analysis'] = json.loads(email['analysis'])
                except:
                    pass
        
        # Cache the result
        server_cache.set(cache_key, emails)
        
        return jsonify({
            "success": True,
            "emails": emails
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/<int:email_id>', methods=['GET'])
@login_required
def get_email(email_id):
    """Get a specific email by ID"""
    try:
        user_id = request.current_user['id']
        email = database.get_email_by_id(email_id, user_id)
        
        if email:
            return jsonify({
                "success": True,
                "email": email
            })
        else:
            return jsonify({
                "success": False,
                "error": "Email not found"
            }), 404
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/<int:email_id>/thread', methods=['GET'])
@login_required
def get_email_thread(email_id):
    """Get conversation thread for an email"""
    try:
        user_id = request.current_user['id']
        thread = database.get_conversation_thread(email_id, user_id)
        
        # Parse analysis JSON for each email in thread
        for email in thread:
            if email.get('analysis'):
                try:
                    email['analysis'] = json.loads(email['analysis'])
                except:
                    pass
        
        return jsonify({
            "success": True,
            "thread": thread,
            "thread_count": len(thread)
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/<int:email_id>/reply', methods=['POST'])
@login_required
def reply_to_email(email_id):
    """Reply to an email - continue conversation"""
    try:
        data = request.json
        user_id = request.current_user.get('id')
        
        # Get user settings for email config
        user_settings = database.get_user_settings(user_id) if user_id else None
        
        # Check if email is configured
        if not user_settings or not user_settings.get('sender_email') or not user_settings.get('sender_password'):
            return jsonify({
                "success": False,
                "error": "Bạn chưa cấu hình Email gửi. Vui lòng vào Hồ sơ để cấu hình trước khi gửi email.",
                "error_code": "EMAIL_NOT_CONFIGURED"
            }), 400
        
        # Get original email (with user_id check for security)
        user_id = request.current_user['id']
        original_email = database.get_email_by_id(email_id, user_id)
        if not original_email:
            return jsonify({
                "success": False,
                "error": "Email gốc không tồn tại hoặc bạn không có quyền truy cập"
            }), 404
        
        purpose = data.get('purpose')
        additional_context = data.get('additional_context')
        custom_subject = data.get('custom_subject')
        custom_body = data.get('custom_body')
        
        if not purpose and not custom_body:
            return jsonify({
                "success": False,
                "error": "Vui lòng nhập mục đích trả lời hoặc nội dung email"
            }), 400
        
        # Get user's AI agent
        user_ai_agent = get_ai_agent_for_user(user_settings)
        
        # Get sender info from user settings
        sender_email = user_settings.get('sender_email')
        sender_name = original_email.get('sender_name', 'Người gửi')
        
        # Generate reply email if no custom content
        if custom_subject and custom_body:
            reply_subject = custom_subject
            reply_body = custom_body
        else:
            # Build context for AI to understand the conversation
            context = f"""
Đây là email trả lời trong cuộc hội thoại.

Email gốc đã gửi:
- Tiêu đề: {original_email.get('subject')}
- Nội dung: {original_email.get('body')}

Phản hồi đã nhận được từ {original_email.get('recipient_name')}:
{original_email.get('response_body', 'Chưa có nội dung phản hồi cụ thể')}

{f'Thông tin thêm: {additional_context}' if additional_context else ''}
"""
            # Generate reply using AI
            generated = user_ai_agent.generate_email(
                sender_name=sender_name,
                recipient_name=original_email.get('recipient_name'),
                recipient_email=original_email.get('recipient_email'),
                purpose=purpose,
                tone='professional',
                language='vi',
                additional_context=context
            )
            
            reply_subject = f"Re: {original_email.get('subject')}"
            reply_body = generated.get('body', '')
        
        # Send the reply email using user's email config
        reply_email_service = EmailService()
        reply_email_service.sender_email = sender_email
        reply_email_service.sender_password = user_settings.get('sender_password')
        reply_email_service.smtp_host = user_settings.get('email_host', 'smtp.gmail.com')
        reply_email_service.smtp_port = int(user_settings.get('email_port', 587))
        
        send_success = reply_email_service.send_email(
            recipient_email=original_email.get('recipient_email'),
            subject=reply_subject,
            body=reply_body
        )
        
        if not send_success:
            return jsonify({
                "success": False,
                "error": "Không thể gửi email trả lời. Vui lòng kiểm tra cấu hình email."
            }), 500
        
        # Save reply to database
        reply_id = database.save_reply_email(
            parent_email_id=email_id,
            user_id=user_id,
            sender_name=sender_name,
            sender_email=sender_email,
            recipient_name=original_email.get('recipient_name'),
            recipient_email=original_email.get('recipient_email'),
            subject=reply_subject,
            body=reply_body,
            purpose=purpose or 'Trả lời email',
            message_id=None  # send_email doesn't return message_id
        )
        
        return jsonify({
            "success": True,
            "message": "Đã gửi email trả lời thành công",
            "email_id": reply_id,
            "subject": reply_subject,
            "body": reply_body
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/<int:email_id>/preview-reply', methods=['POST'])
@login_required
def preview_reply_email(email_id):
    """Preview AI-generated reply email"""
    try:
        data = request.json
        user_id = request.current_user.get('id')
        
        # Get user settings
        user_settings = database.get_user_settings(user_id) if user_id else None
        
        # Get original email (with user_id check for security)
        original_email = database.get_email_by_id(email_id, user_id)
        if not original_email:
            return jsonify({
                "success": False,
                "error": "Email gốc không tồn tại hoặc bạn không có quyền truy cập"
            }), 404
        
        purpose = data.get('purpose')
        additional_context = data.get('additional_context')
        
        if not purpose:
            return jsonify({
                "success": False,
                "error": "Vui lòng nhập mục đích trả lời"
            }), 400
        
        # Get user's AI agent
        user_ai_agent = get_ai_agent_for_user(user_settings)
        
        sender_name = original_email.get('sender_name', 'Người gửi')
        
        # Build context for AI
        context = f"""
Đây là email trả lời trong cuộc hội thoại.

Email gốc đã gửi:
- Tiêu đề: {original_email.get('subject')}
- Nội dung: {original_email.get('body')}

Phản hồi đã nhận được từ {original_email.get('recipient_name')}:
{original_email.get('response_body', 'Chưa có nội dung phản hồi cụ thể')}

{f'Thông tin thêm: {additional_context}' if additional_context else ''}
"""
        
        # Generate reply using AI
        generated = user_ai_agent.generate_email(
            sender_name=sender_name,
            recipient_name=original_email.get('recipient_name'),
            recipient_email=original_email.get('recipient_email'),
            purpose=purpose,
            tone='professional',
            language='vi',
            additional_context=context
        )
        
        return jsonify({
            "success": True,
            "subject": f"Re: {original_email.get('subject')}",
            "body": generated.get('body', '')
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/<int:email_id>', methods=['DELETE'])
@login_required
def delete_email(email_id):
    """Delete a specific email by ID"""
    try:
        user_id = request.current_user['id']
        result = database.delete_email(email_id, user_id)
        if not result:
            return jsonify({
                "success": False,
                "error": "Email not found or access denied"
            }), 404
        
        # Invalidate cache
        server_cache.invalidate(f'emails:{user_id}')
        server_cache.invalidate(f'stats:{user_id}')
        
        return jsonify({
            "success": True,
            "message": "Email deleted successfully"
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/delete-multiple', methods=['POST'])
@login_required
def delete_multiple_emails():
    """Delete multiple emails by IDs"""
    try:
        user_id = request.current_user['id']
        data = request.json
        email_ids = data.get('email_ids', [])
        
        if not email_ids:
            return jsonify({
                "success": False,
                "error": "No email IDs provided"
            }), 400
        
        deleted_count = database.delete_emails(email_ids, user_id)
        
        # Invalidate cache
        server_cache.invalidate(f'emails:{user_id}')
        server_cache.invalidate(f'stats:{user_id}')
        
        return jsonify({
            "success": True,
            "message": f"Deleted {deleted_count} emails",
            "deleted_count": deleted_count
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/emails/delete-all', methods=['DELETE'])
@login_required
def delete_all_emails():
    """Delete all sent emails"""
    try:
        user_id = request.current_user['id']
        deleted_count = database.delete_all_emails(user_id)
        
        # Invalidate cache
        server_cache.invalidate(f'emails:{user_id}')
        server_cache.invalidate(f'stats:{user_id}')
        
        return jsonify({
            "success": True,
            "message": f"Deleted all {deleted_count} emails",
            "deleted_count": deleted_count
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/delete-all', methods=['DELETE'])
@login_required
def delete_all_cv_evaluations():
    """Delete all CV evaluations"""
    try:
        user_id = request.current_user['id']
        deleted_count = database.delete_all_cv_evaluations(user_id)
        
        # Invalidate cache
        server_cache.invalidate(f'cv_list:{user_id}')
        server_cache.invalidate(f'stats:{user_id}')
        
        return jsonify({
            "success": True,
            "message": f"Deleted all {deleted_count} CV evaluations",
            "deleted_count": deleted_count
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/<int:cv_id>', methods=['DELETE'])
@login_required
def delete_cv_evaluation(cv_id):
    """Delete a specific CV evaluation by ID"""
    try:
        user_id = request.current_user['id']
        result = database.delete_cv_evaluation(cv_id, user_id)
        if not result:
            return jsonify({
                "success": False,
                "error": "CV not found or access denied"
            }), 404
        return jsonify({
            "success": True,
            "message": "CV evaluation deleted successfully"
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/delete-multiple', methods=['POST'])
@login_required
def delete_multiple_cv_evaluations():
    """Delete multiple CV evaluations by IDs"""
    try:
        user_id = request.current_user['id']
        data = request.json
        cv_ids = data.get('cv_ids', [])
        
        if not cv_ids:
            return jsonify({
                "success": False,
                "error": "No CV IDs provided"
            }), 400
        
        deleted_count = database.delete_cv_evaluations(cv_ids, user_id)
        return jsonify({
            "success": True,
            "message": f"Deleted {deleted_count} CV evaluations",
            "deleted_count": deleted_count
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/data/delete-all', methods=['DELETE'])
@login_required
def delete_all_data():
    """Delete all emails and CV evaluations"""
    try:
        user_id = request.current_user['id']
        email_count = database.delete_all_emails(user_id)
        cv_count = database.delete_all_cv_evaluations(user_id)
        return jsonify({
            "success": True,
            "message": f"Deleted {email_count} emails and {cv_count} CV evaluations",
            "deleted_emails": email_count,
            "deleted_cvs": cv_count
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/data/stats', methods=['GET'])
@login_required
def get_data_stats():
    """Get statistics about emails and CV evaluations"""
    try:
        user_id = request.current_user['id']
        email_count = database.get_email_count(user_id)
        cv_count = database.get_cv_count(user_id)
        return jsonify({
            "success": True,
            "email_count": email_count,
            "cv_count": cv_count
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/check-responses', methods=['POST'])
@login_required
def check_responses():
    """Check for responses once using user's email settings"""
    try:
        # Get user settings from database
        user_settings = database.get_user_settings(request.current_user['id'])
        
        if user_settings and user_settings.get('sender_email') and user_settings.get('sender_password'):
            # Create email service with user's settings
            user_email_service = EmailService()
            user_email_service.sender_email = user_settings.get('sender_email')
            user_email_service.sender_password = user_settings.get('sender_password')
            user_email_service.imap_host = user_settings.get('imap_host', 'imap.gmail.com')
            user_email_service.imap_port = int(user_settings.get('imap_port', 993))
            
            temp_monitor = EmailMonitor(
                user_email_service,
                ai_agent,
                database,
                notification_callback=notification_callback
            )
        else:
            temp_monitor = EmailMonitor(
                email_service,
                ai_agent,
                database,
                notification_callback=notification_callback
            )
        
        # Check responses and get results
        results = temp_monitor.check_responses_with_details()
        
        return jsonify({
            "success": True,
            "message": "Response check completed",
            "results": results
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/check-imap', methods=['GET'])
@login_required
def check_imap():
    """Check IMAP connection status using user's settings"""
    try:
        # Get user settings from database
        user_settings = database.get_user_settings(request.current_user['id'])
        
        if user_settings and user_settings.get('sender_email') and user_settings.get('sender_password'):
            # Create temporary email service with user's settings
            from services.email_service import EmailService
            temp_email_service = EmailService()
            temp_email_service.sender_email = user_settings.get('sender_email')
            temp_email_service.sender_password = user_settings.get('sender_password')
            
            # Use custom IMAP host/port if configured
            imap_host = user_settings.get('imap_host', 'imap.gmail.com')
            imap_port = user_settings.get('imap_port', 993)
            
            result = temp_email_service.check_imap_connection_with_config(imap_host, imap_port)
        else:
            # Fall back to global email service
            result = email_service.check_imap_connection()
        
        return jsonify(result)
    except Exception as e:
        return jsonify({
            "success": False,
            "message": "Error checking IMAP",
            "details": str(e)
        })


@app.route('/api/monitor/start', methods=['POST'])
@login_required
def start_monitor():
    """Start the email monitoring service"""
    global monitor
    
    try:
        if monitor and monitor._running:
            return jsonify({
                "success": False,
                "message": "Monitor is already running"
            })
        
        monitor = EmailMonitor(
            email_service,
            ai_agent,
            database,
            notification_callback=notification_callback
        )
        monitor.start()
        
        return jsonify({
            "success": True,
            "message": "Monitor started successfully"
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/monitor/stop', methods=['POST'])
@login_required
def stop_monitor():
    """Stop the email monitoring service"""
    global monitor
    
    try:
        if monitor:
            monitor.stop()
            monitor = None
            
        return jsonify({
            "success": True,
            "message": "Monitor stopped successfully"
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/monitor/status', methods=['GET'])
@login_required
def monitor_status():
    """Get monitor status"""
    global monitor
    
    return jsonify({
        "success": True,
        "running": monitor is not None and monitor._running
    })


@app.route('/api/analyze-response', methods=['POST'])
@login_required
def analyze_response():
    """Manually analyze a response"""
    try:
        data = request.json
        
        email_id = data.get('email_id')
        response_text = data.get('response_text')
        response_subject = data.get('response_subject', 'Re: Response')
        
        if not email_id or not response_text:
            return jsonify({
                "success": False,
                "error": "Missing required fields"
            }), 400
        
        processor = ManualResponseProcessor(
            email_service,
            ai_agent,
            database
        )
        
        result = processor.process_response(
            sent_email_id=email_id,
            response_text=response_text,
            response_subject=response_subject
        )
        
        return jsonify({
            "success": True,
            "analysis": result['analysis'],
            "notification": result['notification'],
            "notification_sent": result['notification_sent']
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/events')
def events():
    """Server-Sent Events for real-time updates"""
    def generate():
        while True:
            try:
                event = event_queue.get(timeout=30)
                yield f"data: {json.dumps(event)}\n\n"
            except queue.Empty:
                yield f"data: {json.dumps({'type': 'heartbeat'})}\n\n"
    
    return Response(generate(), mimetype='text/event-stream')


# ==================== CV Evaluation API ====================

@app.route('/api/cv/upload', methods=['POST'])
@login_required
def upload_cv_file():
    """Upload và trích xuất nội dung từ file CV"""
    try:
        if 'file' not in request.files:
            return jsonify({
                "success": False,
                "error": "Không tìm thấy file"
            }), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({
                "success": False,
                "error": "Chưa chọn file"
            }), 400
        
        # Get file extension
        filename = file.filename
        file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        
        # Validate extension
        allowed_extensions = ['pdf', 'doc', 'docx', 'txt']
        if file_ext not in allowed_extensions:
            return jsonify({
                "success": False,
                "error": f"Định dạng file không hỗ trợ. Chỉ chấp nhận: {', '.join(allowed_extensions)}"
            }), 400
        
        # Read file content
        file_content = file.read()
        
        # Extract text based on file type
        extracted_text = ""
        
        if file_ext == 'txt':
            # Plain text file
            try:
                extracted_text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                extracted_text = file_content.decode('latin-1')
                
        elif file_ext == 'pdf':
            # PDF file
            try:
                import PyPDF2
                import io
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                extracted_text = '\n'.join(text_parts)
                
            except ImportError:
                return jsonify({
                    "success": False,
                    "error": "Thư viện PyPDF2 chưa được cài đặt. Chạy: pip install PyPDF2"
                }), 500
            except Exception as e:
                return jsonify({
                    "success": False,
                    "error": f"Lỗi đọc file PDF: {str(e)}"
                }), 500
                
        elif file_ext in ['doc', 'docx']:
            # Word document
            try:
                from docx import Document
                import io
                
                doc = Document(io.BytesIO(file_content))
                text_parts = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            text_parts.append(row_text)
                
                extracted_text = '\n'.join(text_parts)
                
            except ImportError:
                return jsonify({
                    "success": False,
                    "error": "Thư viện python-docx chưa được cài đặt. Chạy: pip install python-docx"
                }), 500
            except Exception as e:
                return jsonify({
                    "success": False,
                    "error": f"Lỗi đọc file Word: {str(e)}"
                }), 500
        
        # Check if we got any content
        if not extracted_text.strip():
            return jsonify({
                "success": False,
                "error": "Không thể trích xuất nội dung từ file. File có thể trống hoặc chứa hình ảnh."
            }), 400
        
        return jsonify({
            "success": True,
            "content": extracted_text,
            "filename": filename,
            "file_type": file_ext
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/evaluate', methods=['POST'])
@login_required
def evaluate_cv():
    """Đánh giá CV ứng viên"""
    try:
        data = request.json
        
        candidate_name = data.get('candidate_name')
        candidate_email = data.get('candidate_email')
        job_title = data.get('job_title')
        job_requirements = data.get('job_requirements')
        cv_content = data.get('cv_content')
        company_name = data.get('company_name', '')
        auto_send = data.get('auto_send', True)  # Tự động gửi email nếu >= 85%
        
        if not all([candidate_name, candidate_email, job_title, job_requirements, cv_content]):
            return jsonify({
                "success": False,
                "error": "Thiếu thông tin bắt buộc"
            }), 400
        
        # Get user settings to check for custom API key
        user_id = request.current_user.get('id')
        user_settings = database.get_user_settings(user_id) if user_id else None
        
        # Use user's CV evaluator if they have custom API key
        user_cv_evaluator = get_cv_evaluator_for_user(user_settings)
        
        # Đánh giá CV
        evaluation = user_cv_evaluator.evaluate_cv(
            cv_content=cv_content,
            job_title=job_title,
            job_requirements=job_requirements,
            company_name=company_name
        )
        
        # Lưu kết quả đánh giá
        cv_id = database.save_cv_evaluation(
            candidate_name=candidate_name,
            candidate_email=candidate_email,
            job_title=job_title,
            job_requirements=job_requirements,
            cv_content=cv_content,
            company_name=company_name,
            overall_score=evaluation.get('overall_score', 0),
            is_qualified=evaluation.get('is_qualified', False),
            evaluation_result=evaluation
        )
        
        result = {
            "success": True,
            "cv_id": cv_id,
            "evaluation": evaluation,
            "is_qualified": evaluation.get('is_qualified', False),
            "overall_score": evaluation.get('overall_score', 0)
        }
        
        # Tự động gửi email nếu đạt >= 85% và auto_send = True
        if auto_send and evaluation.get('is_qualified', False):
            email_result = send_cv_invitation_email(cv_id, evaluation, candidate_name, 
                                                     candidate_email, job_title, company_name)
            result['email_sent'] = email_result.get('success', False)
            result['email_id'] = email_result.get('email_id')
            result['message'] = "Ứng viên đạt yêu cầu! Email mời phỏng vấn đã được gửi."
        elif evaluation.get('is_qualified', False):
            result['message'] = "Ứng viên đạt yêu cầu! Bạn có thể gửi email mời phỏng vấn."
        else:
            result['message'] = f"Ứng viên chưa đạt yêu cầu (Điểm: {evaluation.get('overall_score', 0)}/100, cần >= 85)"
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


def send_cv_invitation_email(cv_id, evaluation, candidate_name, candidate_email, job_title, company_name):
    """Helper function to send invitation email"""
    try:
        # Tạo email mời phỏng vấn
        email_content = cv_evaluator.generate_interview_invitation(
            candidate_name=candidate_name,
            candidate_email=candidate_email,
            job_title=job_title,
            company_name=company_name or "Công ty",
            evaluation_result=evaluation
        )
        
        # Gửi email
        success = email_service.send_email(
            recipient_email=candidate_email,
            subject=email_content['subject'],
            body=email_content['body']
        )
        
        if success:
            # Lưu vào database
            email_id = database.save_sent_email(
                sender_name="HR " + (company_name or ""),
                sender_email=SENDER_EMAIL,
                recipient_name=candidate_name,
                recipient_email=candidate_email,
                subject=email_content['subject'],
                body=email_content['body'],
                purpose=f"Mời phỏng vấn vị trí {job_title}"
            )
            
            # Cập nhật trạng thái CV
            database.update_cv_evaluation_email_sent(cv_id, email_id)
            
            return {"success": True, "email_id": email_id}
        
        return {"success": False, "error": "Không thể gửi email"}
        
    except Exception as e:
        return {"success": False, "error": str(e)}


@app.route('/api/cv/send-email/<int:cv_id>', methods=['POST'])
@login_required
def send_cv_email(cv_id):
    """Gửi/gửi lại email cho ứng viên"""
    try:
        user_id = request.current_user['id']
        cv_data = database.get_cv_evaluation_by_id(cv_id, user_id)
        
        if not cv_data:
            return jsonify({
                "success": False,
                "error": "Không tìm thấy CV hoặc bạn không có quyền truy cập"
            }), 404
        
        data = request.json or {}
        email_type = data.get('email_type', 'invitation')  # invitation or rejection
        
        if email_type == 'invitation':
            email_content = cv_evaluator.generate_interview_invitation(
                candidate_name=cv_data['candidate_name'],
                candidate_email=cv_data['candidate_email'],
                job_title=cv_data['job_title'],
                company_name=cv_data['company_name'] or "Công ty",
                evaluation_result=cv_data.get('evaluation_result', {}),
                interview_details=data.get('interview_details')
            )
        else:
            email_content = cv_evaluator.generate_rejection_email(
                candidate_name=cv_data['candidate_name'],
                job_title=cv_data['job_title'],
                company_name=cv_data['company_name'] or "Công ty",
                evaluation_result=cv_data.get('evaluation_result', {})
            )
        
        # Gửi email
        success = email_service.send_email(
            recipient_email=cv_data['candidate_email'],
            subject=email_content['subject'],
            body=email_content['body']
        )
        
        if success:
            email_id = database.save_sent_email(
                sender_name="HR " + (cv_data['company_name'] or ""),
                sender_email=SENDER_EMAIL,
                recipient_name=cv_data['candidate_name'],
                recipient_email=cv_data['candidate_email'],
                subject=email_content['subject'],
                body=email_content['body'],
                purpose=f"{'Mời phỏng vấn' if email_type == 'invitation' else 'Thông báo kết quả'} - {cv_data['job_title']}"
            )
            
            database.update_cv_evaluation_email_sent(cv_id, email_id)
            
            return jsonify({
                "success": True,
                "email_id": email_id,
                "message": "Email đã được gửi thành công!"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Không thể gửi email"
            }), 500
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/allow-resend/<int:cv_id>', methods=['POST'])
@login_required
def allow_resend_cv_email(cv_id):
    """Cho phép gửi lại email cho CV (sau khi đã có phản hồi)"""
    try:
        database.allow_resend_email(cv_id)
        return jsonify({
            "success": True,
            "message": "Đã cho phép gửi lại email"
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/list', methods=['GET'])
@login_required
def list_cv_evaluations():
    """Lấy danh sách tất cả CV đã đánh giá"""
    try:
        user_id = request.current_user['id']
        evaluations = database.get_all_cv_evaluations(user_id)
        return jsonify({
            "success": True,
            "evaluations": evaluations
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/<int:cv_id>', methods=['GET'])
@login_required
def get_cv_evaluation(cv_id):
    """Lấy chi tiết một CV đánh giá"""
    try:
        user_id = request.current_user['id']
        cv_data = database.get_cv_evaluation_by_id(cv_id, user_id)
        
        if cv_data:
            # Lấy thông tin email liên quan nếu có
            if cv_data.get('sent_email_id'):
                email_data = database.get_email_by_id(cv_data['sent_email_id'], user_id)
                cv_data['email_info'] = email_data
            
            return jsonify({
                "success": True,
                "cv": cv_data
            })
        else:
            return jsonify({
                "success": False,
                "error": "Không tìm thấy CV"
            }), 404
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/cv/preview-email', methods=['POST'])
@login_required
def preview_cv_email():
    """Xem trước email sẽ gửi cho ứng viên"""
    try:
        data = request.json
        
        candidate_name = data.get('candidate_name')
        candidate_email = data.get('candidate_email')
        job_title = data.get('job_title')
        company_name = data.get('company_name', 'Công ty')
        email_type = data.get('email_type', 'invitation')
        evaluation_result = data.get('evaluation_result', {})
        interview_details = data.get('interview_details')
        
        # Get user settings to check for custom API key
        user_id = request.current_user.get('id')
        user_settings = database.get_user_settings(user_id) if user_id else None
        
        # Use user's CV evaluator if they have custom API key
        user_cv_evaluator = get_cv_evaluator_for_user(user_settings)
        
        if email_type == 'invitation':
            email_content = user_cv_evaluator.generate_interview_invitation(
                candidate_name=candidate_name,
                candidate_email=candidate_email,
                job_title=job_title,
                company_name=company_name,
                evaluation_result=evaluation_result,
                interview_details=interview_details
            )
        else:
            email_content = user_cv_evaluator.generate_rejection_email(
                candidate_name=candidate_name,
                job_title=job_title,
                company_name=company_name,
                evaluation_result=evaluation_result
            )
        
        return jsonify({
            "success": True,
            "subject": email_content['subject'],
            "body": email_content['body']
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/system-info', methods=['GET'])
def get_system_info():
    """Get system information including AI provider"""
    from config.settings import AI_PROVIDER, GEMINI_MODEL
    
    return jsonify({
        "success": True,
        "ai_provider": AI_PROVIDER,
        "ai_model": GEMINI_MODEL if AI_PROVIDER == "gemini" else "Azure OpenAI",
        "sender_email": SENDER_EMAIL,
        "realtime_enabled": True
    })


# ==================== Auto Reply API ====================

# Initialize auto-reply service
from services.auto_reply_service import AutoReplyService
auto_reply_service = None

def get_auto_reply_service():
    """Lazy initialization of auto-reply service"""
    global auto_reply_service
    if auto_reply_service is None:
        print("🔄 Initializing Auto-Reply Service...")
        auto_reply_service = AutoReplyService(database, ai_agent, email_service)
        print("✅ Auto-Reply Service initialized")
    return auto_reply_service

# Initialize auto-reply tables on app start
try:
    print("🔄 Creating auto-reply tables if not exist...")
    _service = get_auto_reply_service()
    print("✅ Auto-reply tables ready")
except Exception as e:
    print(f"⚠️ Failed to initialize auto-reply tables: {e}")


@app.route('/api/auto-reply/settings', methods=['GET'])
@login_required
def get_auto_reply_settings():
    """Get auto-reply settings for current user"""
    try:
        user_id = request.current_user['id']
        service = get_auto_reply_service()
        settings = service.get_user_settings(user_id)
        
        return jsonify({
            "success": True,
            "settings": settings or {
                "enabled": False,
                "auto_send_threshold": 0.9,
                "require_confirmation": True,
                "confirmation_timeout_hours": 24,
                "reply_languages": "vi,en",
                "exclude_keywords": "",
                "only_business_hours": False,
                "business_hours_start": 9,
                "business_hours_end": 18
            }
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/settings', methods=['POST'])
@login_required
def save_auto_reply_settings():
    """Save auto-reply settings for current user"""
    try:
        user_id = request.current_user['id']
        data = request.json
        
        service = get_auto_reply_service()
        success = service.save_user_settings(user_id, data)
        
        if success:
            return jsonify({
                "success": True,
                "message": "Đã lưu cài đặt trả lời tự động"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Không thể lưu cài đặt"
            }), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/drafts', methods=['GET'])
@login_required
def get_auto_reply_drafts():
    """Get all auto-reply drafts for current user"""
    try:
        user_id = request.current_user['id']
        status = request.args.get('status')  # pending, sent, rejected, expired
        
        service = get_auto_reply_service()
        drafts = service.get_all_drafts(user_id, status)
        
        return jsonify({
            "success": True,
            "drafts": drafts
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/drafts/<int:draft_id>', methods=['GET'])
@login_required
def get_auto_reply_draft(draft_id):
    """Get a specific draft"""
    try:
        service = get_auto_reply_service()
        draft = service.get_draft_by_id(draft_id)
        
        if not draft:
            return jsonify({"success": False, "error": "Không tìm thấy draft"}), 404
        
        # Check ownership
        if draft['user_id'] != request.current_user['id']:
            return jsonify({"success": False, "error": "Không có quyền truy cập"}), 403
        
        return jsonify({
            "success": True,
            "draft": draft
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/confirm/<token>', methods=['GET'])
def confirm_auto_reply(token):
    """Confirm and send auto-reply (via email link)"""
    try:
        service = get_auto_reply_service()
        result = service.confirm_and_send(token)
        
        # Return HTML page for better UX
        if result['success']:
            return render_template_string(CONFIRMATION_SUCCESS_HTML, message=result['message'])
        else:
            return render_template_string(CONFIRMATION_ERROR_HTML, error=result['error'])
    except Exception as e:
        return render_template_string(CONFIRMATION_ERROR_HTML, error=str(e))


@app.route('/api/auto-reply/reject/<token>', methods=['GET'])
def reject_auto_reply(token):
    """Reject auto-reply (via email link)"""
    try:
        service = get_auto_reply_service()
        result = service.reject_draft(token)
        
        if result['success']:
            return render_template_string(REJECTION_SUCCESS_HTML, message=result['message'])
        else:
            return render_template_string(CONFIRMATION_ERROR_HTML, error=result['error'])
    except Exception as e:
        return render_template_string(CONFIRMATION_ERROR_HTML, error=str(e))


@app.route('/api/auto-reply/confirm/<token>', methods=['POST'])
@login_required
def confirm_auto_reply_api(token):
    """Confirm auto-reply via API (from UI)"""
    try:
        service = get_auto_reply_service()
        result = service.confirm_and_send(token)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/reject/<token>', methods=['POST'])
@login_required
def reject_auto_reply_api(token):
    """Reject auto-reply via API (from UI)"""
    try:
        service = get_auto_reply_service()
        result = service.reject_draft(token)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/drafts/<int:draft_id>', methods=['DELETE'])
@login_required
def delete_auto_reply_draft(draft_id):
    """Delete a specific auto-reply draft"""
    try:
        user_id = request.current_user['id']
        service = get_auto_reply_service()
        result = service.delete_draft(draft_id, user_id)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/auto-reply/drafts', methods=['DELETE'])
@login_required
def delete_all_auto_reply_drafts():
    """Delete all auto-reply drafts for current user"""
    try:
        user_id = request.current_user['id']
        service = get_auto_reply_service()
        result = service.delete_all_drafts(user_id)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== YouTube Analyzer API ====================

@app.route('/api/youtube/analyze', methods=['POST'])
@login_required
def analyze_youtube_channel():
    """Analyze a YouTube channel and estimate earnings"""
    try:
        from services.youtube_analyzer import get_youtube_analyzer
        
        data = request.get_json()
        url = data.get('url', '').strip()
        
        if not url:
            return jsonify({'success': False, 'error': 'Vui lòng nhập link kênh YouTube'})
        
        # Get YouTube API key from environment (optional)
        youtube_api_key = os.environ.get('YOUTUBE_API_KEY')
        
        analyzer = get_youtube_analyzer(youtube_api_key)
        result = analyzer.analyze_channel(url)
        
        return jsonify(result)
    except Exception as e:
        logger.error(f"YouTube analysis error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# HTML Templates for confirmation pages
CONFIRMATION_SUCCESS_HTML = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Email đã gửi thành công</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Segoe UI', sans-serif; 
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 20px;
        }
        .card {
            background: white;
            border-radius: 20px;
            padding: 50px 40px;
            text-align: center;
            max-width: 450px;
            box-shadow: 0 25px 50px rgba(0,0,0,0.15);
        }
        .icon {
            width: 80px;
            height: 80px;
            background: linear-gradient(135deg, #10b981 0%, #059669 100%);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 25px;
            font-size: 40px;
        }
        h1 { color: #047857; margin-bottom: 15px; font-size: 24px; }
        p { color: #6b7280; line-height: 1.6; margin-bottom: 25px; }
        .btn {
            display: inline-block;
            background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
            color: white;
            padding: 14px 32px;
            border-radius: 10px;
            text-decoration: none;
            font-weight: 600;
        }
    </style>
</head>
<body>
    <div class="card">
        <div class="icon">✅</div>
        <h1>Thành công!</h1>
        <p>{{ message }}</p>
        <a href="/" class="btn">Về trang chủ</a>
    </div>
</body>
</html>
"""

CONFIRMATION_ERROR_HTML = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Có lỗi xảy ra</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Segoe UI', sans-serif; 
            background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%);
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 20px;
        }
        .card {
            background: white;
            border-radius: 20px;
            padding: 50px 40px;
            text-align: center;
            max-width: 450px;
            box-shadow: 0 25px 50px rgba(0,0,0,0.15);
        }
        .icon {
            width: 80px;
            height: 80px;
            background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 25px;
            font-size: 40px;
        }
        h1 { color: #dc2626; margin-bottom: 15px; font-size: 24px; }
        p { color: #6b7280; line-height: 1.6; margin-bottom: 25px; }
        .error { background: #fef2f2; color: #991b1b; padding: 15px; border-radius: 10px; margin-bottom: 25px; }
        .btn {
            display: inline-block;
            background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
            color: white;
            padding: 14px 32px;
            border-radius: 10px;
            text-decoration: none;
            font-weight: 600;
        }
    </style>
</head>
<body>
    <div class="card">
        <div class="icon">❌</div>
        <h1>Có lỗi xảy ra</h1>
        <div class="error">{{ error }}</div>
        <a href="/" class="btn">Về trang chủ</a>
    </div>
</body>
</html>
"""

REJECTION_SUCCESS_HTML = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Đã từ chối email</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { 
            font-family: 'Segoe UI', sans-serif; 
            background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            padding: 20px;
        }
        .card {
            background: white;
            border-radius: 20px;
            padding: 50px 40px;
            text-align: center;
            max-width: 450px;
            box-shadow: 0 25px 50px rgba(0,0,0,0.15);
        }
        .icon {
            width: 80px;
            height: 80px;
            background: linear-gradient(135deg, #6b7280 0%, #4b5563 100%);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            margin: 0 auto 25px;
            font-size: 40px;
        }
        h1 { color: #374151; margin-bottom: 15px; font-size: 24px; }
        p { color: #6b7280; line-height: 1.6; margin-bottom: 25px; }
        .btn {
            display: inline-block;
            background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
            color: white;
            padding: 14px 32px;
            border-radius: 10px;
            text-decoration: none;
            font-weight: 600;
        }
    </style>
</head>
<body>
    <div class="card">
        <div class="icon">🚫</div>
        <h1>Đã từ chối</h1>
        <p>{{ message }}</p>
        <a href="/" class="btn">Về trang chủ</a>
    </div>
</body>
</html>
"""


# ==================== Chatbot API ====================

from services.chatbot_service import ChatbotService
chatbot_service = None

def get_chatbot_service():
    """Lazy initialization of chatbot service"""
    global chatbot_service
    if chatbot_service is None:
        chatbot_service = ChatbotService(database, ai_agent)
    return chatbot_service


@app.route('/api/chatbot/query', methods=['POST'])
@login_required
def chatbot_query():
    """Process chatbot query and return response with data"""
    try:
        user_id = request.current_user['id']
        is_admin = request.current_user.get('role') == 'admin'
        data = request.json
        query = data.get('query', '').strip()
        session_id = data.get('session_id')
        
        if not query:
            return jsonify({
                "success": False,
                "error": "Vui lòng nhập câu hỏi"
            }), 400
        
        # Get user settings for AI
        user_settings = database.get_user_settings(user_id)
        
        service = get_chatbot_service()
        result = service.process_query(user_id, query, user_settings, is_admin=is_admin, session_id=session_id)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/stats', methods=['GET'])
@login_required
def chatbot_stats():
    """Get all statistics for chatbot display"""
    try:
        user_id = request.current_user['id']
        is_admin = request.current_user.get('role') == 'admin'
        time_range = request.args.get('time_range')
        
        service = get_chatbot_service()
        email_stats = service.get_email_statistics(user_id, time_range, is_admin=is_admin)
        cv_stats = service.get_cv_statistics(user_id, is_admin=is_admin)
        
        return jsonify({
            "success": True,
            "is_admin": is_admin,
            "email_stats": email_stats,
            "cv_stats": cv_stats
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/chart', methods=['GET'])
@login_required
def chatbot_chart():
    """Get chart data for visualization"""
    try:
        user_id = request.current_user['id']
        is_admin = request.current_user.get('role') == 'admin'
        chart_type = request.args.get('type', 'overview')
        
        service = get_chatbot_service()
        chart_data = service.get_chart_data(user_id, chart_type, is_admin=is_admin)
        
        return jsonify({
            "success": True,
            "is_admin": is_admin,
            "chart": chart_data
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/export', methods=['GET'])
@login_required
def chatbot_export():
    """Export data to Excel file"""
    try:
        user_id = request.current_user['id']
        is_admin = request.current_user.get('role') == 'admin'
        data_type = request.args.get('type', 'all')  # all, emails, cv
        
        service = get_chatbot_service()
        excel_data = service.generate_excel_data(user_id, data_type, is_admin=is_admin)
        
        # Create response with Excel file
        from flask import send_file
        output = io.BytesIO(excel_data)
        output.seek(0)
        
        prefix = "admin_" if is_admin else ""
        filename = f"{prefix}email_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/quick-stats', methods=['GET'])
@login_required
def chatbot_quick_stats():
    """Get quick statistics summary"""
    try:
        user_id = request.current_user['id']
        is_admin = request.current_user.get('role') == 'admin'
        
        service = get_chatbot_service()
        email_stats = service.get_email_statistics(user_id, is_admin=is_admin)
        cv_stats = service.get_cv_statistics(user_id, is_admin=is_admin)
        
        return jsonify({
            "success": True,
            "is_admin": is_admin,
            "stats": {
                "email": {
                    "total_sent": email_stats.get('total_sent', 0),
                    "responded": email_stats.get('responded', 0),
                    "pending": email_stats.get('pending', 0),
                    "response_rate": email_stats.get('response_rate', 0)
                },
                "cv": {
                    "total": cv_stats.get('total', 0),
                    "qualified": cv_stats.get('qualified', 0),
                    "not_qualified": cv_stats.get('not_qualified', 0),
                    "qualification_rate": cv_stats.get('qualification_rate', 0)
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


# ==================== Chat Session APIs ====================

@app.route('/api/chatbot/sessions', methods=['GET'])
@login_required
def get_chat_sessions():
    """Get all chat sessions for current user"""
    try:
        user_id = request.current_user['id']
        sessions = database.get_chat_sessions(user_id)
        return jsonify({
            "success": True,
            "sessions": sessions
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/sessions', methods=['POST'])
@login_required
def create_chat_session():
    """Create a new chat session"""
    try:
        user_id = request.current_user['id']
        data = request.json or {}
        title = data.get('title', 'New Chat')
        
        session_id = database.create_chat_session(user_id, title)
        return jsonify({
            "success": True,
            "session_id": session_id
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/sessions/<int:session_id>', methods=['GET'])
@login_required
def get_chat_session(session_id):
    """Get a specific chat session with messages"""
    try:
        user_id = request.current_user['id']
        session = database.get_chat_session(session_id, user_id)
        
        if not session:
            return jsonify({
                "success": False,
                "error": "Session not found"
            }), 404
        
        messages = database.get_chat_messages(session_id, user_id)
        return jsonify({
            "success": True,
            "session": session,
            "messages": messages
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/sessions/<int:session_id>', methods=['PUT'])
@login_required
def update_chat_session(session_id):
    """Update chat session title"""
    try:
        user_id = request.current_user['id']
        data = request.json or {}
        title = data.get('title')
        
        database.update_chat_session(session_id, user_id, title)
        return jsonify({
            "success": True
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


@app.route('/api/chatbot/sessions/<int:session_id>', methods=['DELETE'])
@login_required
def delete_chat_session(session_id):
    """Delete a chat session"""
    try:
        user_id = request.current_user['id']
        permanent = request.args.get('permanent', 'false').lower() == 'true'
        
        if permanent:
            success = database.delete_chat_session_permanent(session_id, user_id)
        else:
            success = database.delete_chat_session(session_id, user_id)
        
        return jsonify({
            "success": success
        })
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500


if __name__ == '__main__':
    # Sử dụng SocketIO để chạy server với WebSocket support
    print("🚀 Starting Email AI Agent with Realtime WebSocket support...")
    print(f"🤖 AI Provider: {AI_PROVIDER.upper()}")
    print("📡 WebSocket enabled for instant notifications")
    
    # Auto-start monitor if configured
    if AUTO_START_MONITOR:
        print("🔄 Auto-starting email monitor...")
        auto_start_monitor()
    else:
        print("ℹ️ Monitor will start on first WebSocket connection")
    
    socketio.run(app, debug=True, port=5000, allow_unsafe_werkzeug=True)
